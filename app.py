import os
import time
import openai
from flask import Flask, render_template, request, jsonify, send_file
from dotenv import load_dotenv
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
import io
from fpdf import FPDF
import re
import logging
import json

# Configure logging
logging.basicConfig(level=logging.INFO)

# Load environment variables
load_dotenv()

# Initialize Flask app
app = Flask(__name__)

# Initialize OpenAI clients for multiple agents
class StoryAgents:
    def __init__(self):
        self.plot_architect = openai.OpenAI(api_key=os.getenv("OPENAI_API_KEY"))
        self.narrative_developer = openai.OpenAI(api_key=os.getenv("OPENAI_API_KEY"))
        self.dialogue_enhancer = openai.OpenAI(api_key=os.getenv("OPENAI_API_KEY"))
        self.continuity_expert = openai.OpenAI(api_key=os.getenv("OPENAI_API_KEY"))

    def generate_initial_plot_outline(self, title, description, num_chapters):
        """Generate a comprehensive plot outline for the entire story"""
        response = self.plot_architect.chat.completions.create(
            model="gpt-3.5-turbo",
            messages=[
                {
                    "role": "system", 
                    "content": """You are a master storyteller and plot architect. 
                    Create a comprehensive story structure that ensures narrative cohesion, 
                    character development, and engaging plot progression."""
                },
                {
                    "role": "user", 
                    "content": f"""Develop a detailed plot outline for a story with the following parameters:
                    - Title: {title}
                    - Description: {description}
                    - Number of Chapters: {num_chapters}

                    For each chapter, provide:
                    1. Key plot points
                    2. Character arcs
                    3. Emotional trajectory
                    4. Potential conflicts and resolutions
                    5. Thematic development

                    Ensure:
                    - Clear narrative arc
                    - Consistent character motivations
                    - Gradual plot escalation
                    - Meaningful character transformations"""
                }
            ],
            temperature=0.7,
            max_tokens=2000
        )
        return response.choices[0].message.content

    def generate_chapter(self, previous_chapters, plot_outline, chapter_number):
        """Generate a chapter with context from previous chapters"""
        context = "\n\n".join([chapter['content'] for chapter in previous_chapters]) if previous_chapters else ""
        
        response = self.narrative_developer.chat.completions.create(
            model="gpt-3.5-turbo",
            messages=[
                {
                    "role": "system", 
                    "content": """You are an expert novelist specializing in crafting compelling narrative chapters. 
                    Ensure narrative flow, character depth, and engaging storytelling."""
                },
                {
                    "role": "user", 
                    "content": f"""Generate Chapter {chapter_number}

                    Previous Story Context:
                    {context}

                    Plot Outline for Context:
                    {plot_outline}

                    Chapter Generation Guidelines:
                    1. Create a unique, intriguing chapter title
                    2. Write 1200-1500 words
                    3. Advance the plot meaningfully
                    4. Develop characters
                    5. Maintain consistent tone and style"""
                }
            ],
            temperature=0.8,
            max_tokens=2000
        )
        
        chapter_content = response.choices[0].message.content
        
        # Extract chapter title (assumed to be first line)
        chapter_lines = chapter_content.strip().split('\n')
        chapter_title = chapter_lines[0].strip()

        # Clean up chapter title - remove any formatting and "Chapter X:" prefixes
        chapter_title = re.sub(r'\*\*|\*|\"\"|\"|#', '', chapter_title)
        chapter_title = re.sub(r'^Chapter\s+\d+\s*:?\s*', '', chapter_title).strip()
        
        # Remove title from content
        chapter_body = '\n'.join(chapter_lines[1:]).strip()

        # Clean up any remaining markdown formatting in the chapter body
        chapter_body = re.sub(r'\*\*|\*|#', '', chapter_body).strip()
        
        return {
            "number": chapter_number,
            "title": chapter_title,
            "content": chapter_body
        }

def generate_story(title, description, num_chapters):
    """Enhanced story generation with multi-agent approach"""
    agents = StoryAgents()
    
    # Generate comprehensive plot outline
    plot_outline = agents.generate_initial_plot_outline(title, description, num_chapters)
    
    # Initialize story structure
    story = {
        "title": title,
        "description": description,
        "plot_outline": plot_outline,
        "chapters": []
    }
    
    # Generate title and subtitle
    title_response = agents.plot_architect.chat.completions.create(
        model="gpt-3.5-turbo",
        messages=[
            {"role": "system", "content": "Create professional, engaging book titles and subtitles."},
            {"role": "user", "content": f"Generate a compelling book title and subtitle for a story about: {description}"}
        ],
        temperature=0.7,
        max_tokens=100
    )

    improved_title = title_response.choices[0].message.content.strip().strip('"')
    # Clean up any remaining formatting markers
    improved_title = re.sub(r'\*\*|\*|#|Title:\s*', '', improved_title)
    # Only use the improved title if it's valid and not too long
    if improved_title and len(improved_title) <= 100:
        story["original_title"] = title
        title = improved_title
    
    # Blurb generation
    blurb_response = agents.plot_architect.chat.completions.create(
        model="gpt-3.5-turbo",
        messages=[
            {"role": "system", "content": "Create engaging book blurbs that capture the essence of the story."},
            {"role": "user", "content": f"Write a compelling 200-word book blurb for a story titled '{title}' about: {description}"}
        ],
        temperature=0.7,
        max_tokens=300
    )

    blurb = blurb_response.choices[0].message.content.strip()
    # Clean up any markdown formatting
    blurb = re.sub(r'\*\*|\*|#', '', blurb)
    
    # Parse responses
    story['title'] = title
    story['blurb'] = blurb
    
    # Generate chapters iteratively
    previous_chapters = []
    for chapter_num in range(1, num_chapters + 1):
        chapter = agents.generate_chapter(previous_chapters, plot_outline, chapter_num)
        story['chapters'].append(chapter)
        previous_chapters.append(chapter)
    
    return story

def create_pdf(story):
    try:
        class StoryPDF(FPDF):
            def __init__(self):
                super().__init__()
                self.set_auto_page_break(auto=True, margin=20)
                self.set_margins(25, 25, 25)
                self.set_title(story["title"])
                self.chapter_start = False
                self.page_num_offset = 3  # Cover, blurb, TOC pages
                self.set_font('Times', '')  # For regular text
                self.set_font('Times', 'B')  # For bold text
                self.set_font('Times', 'I')  # For italic text
                
            def header(self):
                if self.page_no() > self.page_num_offset:
                    if not self.chapter_start:
                        self.set_font('Times', 'I', 10)
                        self.set_y(10)
                        # Alternate header placement
                        if self.page_no() % 2 == 0:  # Even page
                            self.cell(0, 10, story["title"], 0, 0, 'L')
                        else:  # Odd page
                            self.cell(0, 10, "AI Story Generator", 0, 0, 'R')
                    self.chapter_start = False
            
            def footer(self):
                if self.page_no() > self.page_num_offset:
                    self.set_y(-15)
                    self.set_font('Times', 'I', 10)
                    # Centered page numbers
                    self.cell(0, 10, str(self.page_no() - self.page_num_offset), 0, 0, 'C')
        
        # Initialize PDF
        pdf = StoryPDF()
        
        # Clean title function
        def clean_title(title):
            cleaned = re.sub(r'\*\*|\*|#', '', title)
            cleaned = re.sub(r'^Title:\s*', '', cleaned)
            cleaned = re.sub(r'^Chapter\s*\d+\s*:?\s*', '', cleaned)
            return cleaned.strip()
        
        # Sanitize text for PDF
        def sanitize_for_pdf(text):
            replacements = {
                '\u2014': '-',  # em dash
                '\u2013': '-',  # en dash
                '\u2018': "'",  # left single quote
                '\u2019': "'",  # right single quote
                '\u201C': '"',  # left double quote
                '\u201D': '"',  # right double quote
                '\u2026': '...', # ellipsis
                '\u00A0': ' '   # non-breaking space
            }
            
            for char, replacement in replacements.items():
                text = text.replace(char, replacement)
                
            return text
        
        # Cover Page
        pdf.add_page()
        pdf.set_fill_color(240, 240, 245)
        pdf.rect(0, 60, 210, 100, 'F')
        pdf.ln(80)
        pdf.set_font('Times', 'B', 26)
        
        clean_story_title = sanitize_for_pdf(clean_title(story["title"]))
        
        # Title formatting
        title_words = clean_story_title.split()
        if len(title_words) > 4:
            chunks = []
            current_chunk = []
            current_length = 0
            
            for word in title_words:
                if current_length + len(word) > 20:
                    chunks.append(' '.join(current_chunk))
                    current_chunk = [word]
                    current_length = len(word)
                else:
                    current_chunk.append(word)
                    current_length += len(word) + 1
            
            if current_chunk:
                chunks.append(' '.join(current_chunk))
            
            for i, chunk in enumerate(chunks):
                pdf.cell(0, 16, chunk, 0, 1, 'C')
                if i < len(chunks) - 1:
                    pdf.ln(2)
        else:
            pdf.cell(0, 20, clean_story_title, 0, 1, 'C')
        
        if "subtitle" in story:
            pdf.set_font('Times', 'I', 18)
            clean_subtitle = sanitize_for_pdf(clean_title(story["subtitle"]))
            pdf.cell(0, 15, clean_subtitle, 0, 1, 'C')
            
        pdf.ln(50)
        pdf.set_font('Times', '', 14)
        pdf.cell(0, 10, "Generated by AI Story Generator", 0, 1, 'C')
        
        # Blurb Page
        pdf.add_page()
        pdf.set_font('Times', 'B', 16)
        pdf.cell(0, 20, "About This Book", 0, 1, 'C')
        pdf.ln(10)
        pdf.set_font('Times', '', 12)
        
        if "blurb" in story:
            sanitized_blurb = sanitize_for_pdf(story["blurb"])
            blurb_paragraphs = sanitized_blurb.split('\n\n')
            for para in blurb_paragraphs:
                if para.strip():
                    pdf.multi_cell(0, 6, para.strip())
                    pdf.ln(4)
        
        # Table of Contents
        pdf.add_page()
        pdf.set_font('Times', 'B', 18)
        pdf.cell(0, 20, "Contents", 0, 1, 'C')
        pdf.ln(10)
        pdf.set_font('Times', '', 12)
        
        for chapter in story["chapters"]:
            clean_chapter_title = sanitize_for_pdf(clean_title(chapter['title']))
            chapter['title'] = clean_chapter_title
            
            chapter_text = f"Chapter {chapter['number']}: {clean_chapter_title}"
            text_width = pdf.get_string_width(chapter_text)
            page_text = str(chapter['number'] + 3)
            page_width = pdf.get_string_width(page_text)
            
            available_width = pdf.w - 50 - text_width - page_width
            
            pdf.cell(text_width + 5, 8, chapter_text, 0, 0)
            
            dot_width = pdf.get_string_width('.')
            num_dots = int(available_width / dot_width)
            dots = '.' * num_dots
            pdf.set_font('Times', '', 10)
            pdf.cell(available_width, 8, dots, 0, 0, 'C')
            
            pdf.set_font('Times', '', 12)
            pdf.cell(page_width, 8, page_text, 0, 1, 'R')
        
        # Chapters
        for chapter in story["chapters"]:
            pdf.chapter_start = True
            pdf.add_page()
            
            pdf.set_font('Times', 'B', 18)
            pdf.cell(0, 15, f"Chapter {chapter['number']}", 0, 1, 'C')
            
            pdf.set_font('Times', 'B', 16)
            pdf.cell(0, 10, chapter['title'], 0, 1, 'C')
            
            pdf.line(pdf.w/4, pdf.y + 5, 3*pdf.w/4, pdf.y + 5)
            pdf.ln(15)
            
            sanitized_content = sanitize_for_pdf(chapter['content'])
            paragraphs = re.split(r'\n\n+', sanitized_content)
            
            pdf.set_font('Times', '', 12)
            
            for paragraph in paragraphs:
                paragraph = paragraph.strip()
                if not paragraph:
                    continue
                
                if paragraph.startswith('"') or paragraph.startswith('"'):
                    pdf.multi_cell(0, 6, paragraph)
                elif paragraph.strip() == "* * *" or paragraph.strip() == "***":
                    pdf.ln(4)
                    pdf.set_font('Times', 'B', 12)
                    pdf.cell(0, 6, "* * *", 0, 1, 'C')
                    pdf.ln(4)
                    pdf.set_font('Times', '', 12)
                else:
                    pdf.set_left_margin(pdf.l_margin + 10)
                    pdf.multi_cell(0, 6, paragraph, align='J')
                    pdf.set_left_margin(pdf.l_margin - 10)
                
                pdf.ln(4)
        
        # End Page
        pdf.add_page()
        pdf.set_y(pdf.h / 2 - 10)
        pdf.set_font('Times', 'B', 14)
        pdf.cell(0, 10, "The End", 0, 1, 'C')
        pdf.ln(10)
        pdf.set_font('Times', 'I', 16)
        pdf.cell(0, 10, "* * *", 0, 1, 'C')
        
        # Save to buffer
        buffer = io.BytesIO()
        pdf_output = pdf.output(dest='S').encode('latin-1', errors='replace')
        buffer.write(pdf_output)
        buffer.seek(0)
        return buffer
        
    except Exception as e:
        import traceback
        logging.error(f"PDF generation failed: {str(e)}")
        logging.error(traceback.format_exc())
        raise

def create_docx(story):
    doc = Document()
    
    # Set document margins and page setup
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
        section.different_first_page_header_footer = True
        section.header.is_linked_to_previous = False
        section.footer.is_linked_to_previous = False
    
    # Custom header for pages after first page
    for section in sections:
        first_page_header = section.first_page_header
        first_page_header.is_linked_to_previous = False
        
        default_header = section.header
        default_header.is_linked_to_previous = False
        
        if default_header:
            for i, paragraph in enumerate(default_header.paragraphs):
                paragraph.text = story["title"] if i % 2 == 0 else "AI Story Generator"
                paragraph.alignment = (WD_ALIGN_PARAGRAPH.LEFT 
                                       if i % 2 == 0 
                                       else WD_ALIGN_PARAGRAPH.RIGHT)
    
    # Cover page
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run(story["title"])
    title_run.font.size = Pt(24)
    title_run.font.bold = True
    
    # Subtitle with improved styling
    if "subtitle" in story and story["subtitle"]:
        subtitle_para = doc.add_paragraph()
        subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle_run = subtitle_para.add_run(story["subtitle"])
        subtitle_run.italic = True
        subtitle_run.font.size = Pt(16)
    
    # Add page break after cover
    doc.add_page_break()
    
    # Blurb page with improved formatting
    if "blurb" in story:
        blurb_heading = doc.add_heading("About This Book", 1)
        blurb_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        blurb_para = doc.add_paragraph()
        blurb_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        blurb_para.add_run(story["blurb"])
        doc.add_page_break()
    
    # Table of Contents with improved formatting
    toc_heading = doc.add_heading("Contents", 1)
    toc_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    for idx, chapter in enumerate(story["chapters"]):
        toc_entry = doc.add_paragraph()
        toc_entry.add_run(f"Chapter {chapter['number']}: {chapter['title']}")
        
        toc_entry.add_run("\t")
        toc_entry.add_run(str(idx + 4))
        
        paragraph_format = toc_entry.paragraph_format
        paragraph_format.tab_stops.add_tab_stop(Inches(6), alignment=WD_PARAGRAPH_ALIGNMENT.RIGHT)
    
    # Add page break after TOC
    doc.add_page_break()
    
    # Chapter generation
    for chapter in story["chapters"]:
        chapter_heading = doc.add_heading(f"Chapter {chapter['number']}", 1)
        chapter_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_para.add_run(chapter['title'])
        title_run.bold = True
        title_run.font.size = Pt(14)
        
        line_para = doc.add_paragraph()
        line_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        line_para.add_run("__________________")
        
        paragraphs = chapter["content"].split('\n\n')
        for paragraph in paragraphs:
            paragraph = paragraph.strip()
            if not paragraph:
                continue
            
            p = doc.add_paragraph()
            
            if paragraph.startswith('"') or paragraph.startswith('"'):
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                p.add_run(paragraph)
            
            elif paragraph.strip() == "* * *" or paragraph.strip() == "***":
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.add_run("* * *")
                p.paragraph_format.space_before = Pt(12)
                p.paragraph_format.space_after = Pt(12)
            
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p.paragraph_format.first_line_indent = Inches(0.3)
                p.add_run(paragraph)
            
            for run in p.runs:
                run.font.size = Pt(12)
            
            p.paragraph_format.space_after = Pt(6)
        
        # Page break between chapters
        if chapter != story["chapters"][-1]:
            doc.add_page_break()
    
    # End page
    doc.add_page_break()
    end_para = doc.add_paragraph()
    end_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    end_para.space_before = Pt(200)
    end_run = end_para.add_run("The End")
    end_run.bold = True
    end_run.font.size = Pt(16)
    
    # Save to buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    return buffer

# Routes
@app.route('/')
def index():
    return render_template('index.html')

@app.route('/generate', methods=['POST'])
def generate():
    data = request.json
    title = data.get('title', '').strip()
    description = data.get('description', '').strip()
    num_chapters = int(data.get('num_chapters', 3))
    
    # Validate inputs
    if not title:
        return jsonify({"status": "error", "message": "Title is required"})
    if not description:
        return jsonify({"status": "error", "message": "Description is required"})
    if num_chapters < 1 or num_chapters > 10:
        return jsonify({"status": "error", "message": "Number of chapters must be between 1 and 10"})
    
    try:
        story = generate_story(title, description, num_chapters)
        return jsonify({"status": "success", "story": story})
    except Exception as e:
        logging.error(f"Story generation error: {str(e)}")
        return jsonify({"status": "error", "message": str(e)})

@app.route('/download', methods=['POST'])
def download():
    try:
        data = request.get_json()
        if not data:
            return jsonify({"status": "error", "message": "No JSON data received"}), 400
            
        story = data.get('story')
        format_type = data.get('format', 'pdf')
        
        if not story:
            return jsonify({"status": "error", "message": "No story data provided"}), 400
        
        if format_type == 'pdf':
            try:
                buffer = create_pdf(story)
                
                return send_file(
                    buffer,
                    mimetype='application/pdf',
                    as_attachment=True,
                    download_name=f"{story['title'].replace(' ', '_')}.pdf"
                )
            except Exception as e:
                import traceback
                error_details = traceback.format_exc()
                logging.error(f"PDF generation failed: {error_details}")
                return jsonify({"status": "error", "message": f"PDF generation failed: {str(e)}"}), 500
                
        elif format_type == 'docx':
            try:
                buffer = create_docx(story)
                
                return send_file(
                    buffer,
                    mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                    as_attachment=True,
                    download_name=f"{story['title'].replace(' ', '_')}.docx"
                )
            except Exception as e:
                import traceback
                error_details = traceback.format_exc()
                logging.error(f"DOCX generation failed: {error_details}")
                return jsonify({"status": "error", "message": f"DOCX generation failed: {str(e)}"}), 500
                
        else:
            return jsonify({"status": "error", "message": "Invalid format specified"}), 400
            
    except Exception as e:
        import traceback
        error_details = traceback.format_exc()
        logging.error(f"Download route error: {error_details}")
        return jsonify({"status": "error", "message": str(e), "details": error_details}), 500
    
if __name__ == '__main__':
    app.run(debug=True)
